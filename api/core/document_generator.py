from typing import List, Dict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import openai
import re
from .config import settings

openai.api_key = settings.OPENAI_API_KEY

class DocumentGenerator:
    @staticmethod
    def generate_summary(messages: List[Dict[str,str]]) -> str:
        """
            This method generates a summary from the conversation messages.
            
            Args :
                messages (List[Dict[str, str]]) : The conversation messages.
                
            Returns :
                str : The generated summary.
        """
        conversation_text = "\n".join([f"{message['role']}: {message['content']}" for message in messages])

        summary_prompt = {
            "role": "system",
            "content": f"""Create a concise but comprehensive Web3 x Regenerative Future summary. Keep each section focused and specific:

            ## Executive Summary
            [One paragraph overview of key innovations]

            ### Technical Innovations
            - Core mechanisms proposed
            - Novel combinations discovered
            - Technical challenges addressed

            ### Implementation Framework
            - Key milestones
            - Technical requirements
            - Resource needs

            ### Impact Metrics
            - Environmental KPIs
            - Social impact measures
            - Economic sustainability indicators
            Keep all content specific and actionable. Avoid generic statements.

            # Start of conversation:
            {conversation_text}
            """
        }

        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[summary_prompt],
            temperature=0.7,
            max_tokens=1000,
            stream=False
        )
        return response.choices[0].message.content

    @staticmethod
    def create_document( summary: str,goal: str, messages: List[Dict[str, str]]) -> BytesIO:
        """
        Creates a Word document from the provided input data and returns it as a BytesIO object.

        Args:
            goal (str): The innovation goal.
            summary (str): The generated summary in markdown format.
            messages (List[Dict[str, str]]): The conversation messages.

        Returns:
            BytesIO: The Word document as a binary stream.
        """
        doc = Document()

        title = doc.add_heading('Summarized Document', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Goal:', level=1)
        doc.add_paragraph(goal)

        doc.add_heading('Summary:', level=1)

        def add_paragraph_with_bold_text(text: str, doc):
            """Helper function to process and add text with bold formatting."""
            bold_split = re.split(r'(\*\*.*?\*\*)', text)
            para = doc.add_paragraph()
            for segment in bold_split:
                if segment.startswith('**') and segment.endswith('**'):
                    bold_run = para.add_run(segment[2:-2])
                    bold_run.bold = True
                else:
                    para.add_run(segment)

        summary_lines = summary.strip().split('\n')
        for line in summary_lines:
            line = line.strip()
            if line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            elif re.match(r'^#+\s', line):
                heading_level = line.count('#')
                doc.add_heading(line.lstrip('#').strip(), level=min(heading_level, 4))
            else:
                add_paragraph_with_bold_text(line, doc)

        doc.add_heading('Messages:', level=1)
        for message in messages:
            role_name = message['role']
            doc.add_heading(f"Message from {role_name}:", level=2)

            content_lines = message['content'].strip().split('\n')
            for line in content_lines:
                add_paragraph_with_bold_text(line, doc)

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io