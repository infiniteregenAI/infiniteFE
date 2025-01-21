from typing import List, Dict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import openai
import re

class DocumentGenerator:
    @staticmethod
    def generate_summary(messages: List[Dict[str, str]]) -> str:
        """
        This method generates a detailed and comprehensive summary of the conversation messages.
        
        Args:
            messages (List[Dict[str, str]]): The conversation messages.
        
        Returns:
            str: The generated summary.
        """
        conversation_text = "\n".join([f"{message['role']}: {message['content']}" for message in messages])

        summary_prompt = {
            "role": "system",
            "content": f"""Create a very detailed and clear summary of the following conversation. Ensure the summary:
            - Captures all key points and ideas discussed.
            - Includes specific details, avoiding generic statements.
            - Maintains coherence and structure for easy readability.

            Format the summary as follows:
            ## Overview
            Provide a one-paragraph high-level overview of the conversation.

            ## Key Points
            List the key points discussed, categorized where applicable.

            ## Additional Insights
            Highlight additional observations, implications, or conclusions from the conversation.

            # Start of Conversation:
            {conversation_text}
            """
        }

        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[summary_prompt],
            temperature=0.6,
            max_tokens=16000,
            stream=False
        )
        return response.choices[0].message.content

    @staticmethod
    def create_document(summary: str, goal: str, messages: List[Dict[str, str]]) -> BytesIO:
        """
        Creates a Word document from the provided input data and returns it as a BytesIO object.

        Args:
            goal (str): The innovation goal.
            summary (str): The generated summary in markdown format.
            messages (List[Dict[str, str]]): The conversation messages.

        Returns:
            BytesIO: The Word document as a binary stream.
        """
        doc = Document()

        title = doc.add_heading('Summarized Document', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Goal:', level=1)
        doc.add_paragraph(goal)

        doc.add_heading('Summary:', level=1)

        def add_bold_text_to_paragraph(paragraph, text: str):
            """Helper function to add text with bold formatting to a paragraph."""
            bold_split = re.split(r'(\*\*.*?\*\*)', text)
            for segment in bold_split:
                if segment.startswith('**') and segment.endswith('**'):
                    bold_run = paragraph.add_run(segment[2:-2])
                    bold_run.bold = True
                else:
                    paragraph.add_run(segment)

        def add_heading_with_bold_text(doc, text: str, level: int):
            """Helper function to add a heading with bold text handling."""
            text = text.strip()
            bold_split = re.split(r'(\*\*.*?\*\*)', text)
            heading = doc.add_heading(level=level)
            for segment in bold_split:
                if segment.startswith('**') and segment.endswith('**'):
                    bold_run = heading.add_run(segment[2:-2])
                    bold_run.bold = True
                else:
                    heading.add_run(segment)

        summary_lines = summary.strip().split('\n')
        for line in summary_lines:
            line = line.strip()
            if re.match(r'^(\*|-)\s', line):
                para = doc.add_paragraph(style='List Bullet')
                add_bold_text_to_paragraph(para, line[2:].strip())
            elif re.match(r'^#+\s', line):
                heading_level = line.count('#')
                add_heading_with_bold_text(doc, line.lstrip('#').strip(), level=min(heading_level, 4))
            else:
                para = doc.add_paragraph()
                add_bold_text_to_paragraph(para, line)

        doc.add_heading('Messages:', level=1)
        for message in messages:
            role_name = message['role']
            doc.add_heading(f"Message from {role_name}:", level=2)

            content_lines = message['content'].strip().split('\n')
            for line in content_lines:
                line = line.strip()
                if re.match(r'^(\*|-)\s', line):
                    para = doc.add_paragraph(style='List Bullet')
                    add_bold_text_to_paragraph(para, line[2:].strip())
                elif re.match(r'^#+\s', line):
                    heading_level = line.count('#')
                    add_heading_with_bold_text(doc, line.lstrip('#').strip(), level=min(heading_level, 4))
                else:
                    para = doc.add_paragraph()
                    add_bold_text_to_paragraph(para, line)

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io

